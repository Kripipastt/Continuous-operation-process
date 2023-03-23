import docx
import PyPDF2


def collection(file, format):
    if format == "docx":
        file = docx.Document(file)
        line = "\n".join([file.paragraphs[i].text for i in range(len(file.paragraphs))])
    elif format == "pdf":
        file = PyPDF2.PdfReader(file)
        line = "\n".join([" ".join(file.pages[i].extractText().split("\n")) for i in range(len(file.pages))])
    elif format == "txt":
        with open(file, encoding="UTF-8", mode="r") as file:
            line = file.read()
    return line


if __name__ == "__main__":
    print(collection(r"text_files\Города и годы.txt", "txt"))
